from docx import Document
import io

def build_resume_docx(resume_data, selected_projects, template_path):
    doc = Document(template_path)

    # Replace placeholders in the document
    for para in doc.paragraphs:
        if "{{TITLE}}" in para.text:
            para.text = para.text.replace("{{TITLE}}", resume_data.get("title", ""))
        if "{{SUMMARY}}" in para.text:
            para.text = para.text.replace("{{SUMMARY}}", resume_data.get("summary", ""))
        if "{{SKILLS}}" in para.text:
            para.text = para.text.replace("{{SKILLS}}", ", ".join(resume_data.get("skills", [])))

    # Replace the {{PROJECTS}} section with formatted project entries
    for i, para in enumerate(doc.paragraphs):
        if "{{PROJECTS}}" in para.text:
            para.text = ""  # Clear the placeholder

            for p in selected_projects:
                # Title Line
                title_line = f"{p['title']} ({p.get('client', '-')}, {p.get('duration', '-')})"
                title_para = doc.paragraphs[i].insert_paragraph_before()
                run = title_para.add_run(f"• {title_line}")
                run.bold = True

                # Responsibilities
                if p.get("responsibilities"):
                    for r in p["responsibilities"]:
                        resp_para = doc.paragraphs[i].insert_paragraph_before()
                        resp_para.add_run(f"    - {r}")

                # URL
                if p.get("url"):
                    url_para = doc.paragraphs[i].insert_paragraph_before()
                    url_para.add_run(f"    URL: {p['url']}")

            break  # Only handle the first {{PROJECTS}} placeholder

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def build_cover_letter_docx(cover_letter_text):
    doc = Document()
    for para in cover_letter_text.strip().split("\n\n"):
        doc.add_paragraph(para.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
