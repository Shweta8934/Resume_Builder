from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import os


def fill_docx_template(resume_data, template_name="Resume_Template.docx"):
   
    base_path = os.getcwd()
    template_path = os.path.join(base_path, "..", "templates", template_name)
    doc = Document(template_name)

    replacements = {
        "{{name}}": resume_data.get("name", ""),
        "{{position_title}}": resume_data.get("position_title", ""),
        "{{professional_summary}}": resume_data.get("professional_summary", ""),
        "{{professional_experience}}": "\n".join(
            [f"• {line}" for line in resume_data.get("professional_experience", {}).get("description", [])]
        ),
        "{{technical_skills}}": format_technical_skills(resume_data.get("technical_skills", {})),
        "{{employment_history}}": format_employment_history(resume_data.get("employment_history", [])),
    }

    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text:
                para.text = ""
                run = para.add_run(val)

                if key == "{{name}}":
                    run.bold = True
                    run.font.size = Pt(22)
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                elif key == "{{position_title}}":
                    run.bold = True
                    run.italic = True
                    run.font.size = Pt(16)
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                elif key in ["{{professional_experience}}", "{{employment_history}}"]:
                    run.font.size = Pt(11)
                    para.paragraph_format.left_indent = Inches(0.25)
                    para.paragraph_format.space_after = Pt(4)

                elif key == "{{technical_skills}}":
                    run.font.size = Pt(11)

                else:
                    run.font.size = Pt(11)

        # 👇 Custom handling for project experience
        if "{{project_experience}}" in para.text:
            para.text = ""
            for p in resume_data.get("project_experience", []):
                def add_line(label, value):
                    if value:
                        line = para.insert_paragraph_before()
                        run_label = line.add_run(f"{label}: ")
                        run_label.bold = True
                        run_value = line.add_run(str(value))
                        line.paragraph_format.space_after = Pt(2)

                add_line("Project", p.get("title"))
                add_line("Client", p.get("client"))
                add_line("Duration", p.get("duration"))
                add_line("Role", p.get("role", "N/A"))
                add_line("URL", p.get("url"))
                if p.get("technologies"):
                    add_line("Technologies", ", ".join(p["technologies"]))
                if p.get("key_features"):
                    add_line("Key Features", p["key_features"])
                add_line("Description", p.get("description"))

                if p.get("responsibilities"):
                    resp_title = para.insert_paragraph_before("Responsibilities:")
                    resp_title.runs[0].bold = True
                    for r in p["responsibilities"]:
                        bullet = para.insert_paragraph_before(f"• {r}")
                        bullet.paragraph_format.left_indent = Inches(0.25)

                para.insert_paragraph_before("")  # spacing

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def format_technical_skills(skills):
    result = []
    for category, items in skills.items():
        if items:
            line = f"{category.capitalize()}: {', '.join(items)}"
            result.append(line)
    return "\n".join(result)


def format_employment_history(history):
    result = []
    for job in history:
        result.append(f"{job['company']} - {job['location']}")
        for pos in job.get("positions", []):
            result.append(f"  {pos['title']} ({pos['duration']})")
            for r in pos.get("responsibilities", []):
                result.append(f"    • {r}")
        result.append("")
    return "\n".join(result)
